"""File processing utilities: OCR, PDF, DOCX, and audio transcription."""

import asyncio
import base64
import httpx
import os
import tempfile
from io import BytesIO
from pathlib import Path

from pypdf import PdfReader
from pypdf.errors import PdfReadError
from docx import Document
from PIL import Image

from schemas import FileAttachment, AttachmentType
from logger import get_logger, ocr_logger
from settings import settings

doc_logger = get_logger("DOC")

# Threshold for large document warning (chars)
LARGE_DOCUMENT_THRESHOLD = 6000
# Maximum images to OCR per document (to avoid API rate limits)
MAX_IMAGES_PER_DOCUMENT = 5
# Minimum image dimension to consider for OCR (skip icons)
MIN_IMAGE_DIMENSION = 100


# ── OCR ────────────────────────────────────────────────────────────────────────
async def process_image_ocr(image_data_url: str, image_name: str, language: str = "auto") -> str:
    """Extract text from an image using OCR.space API.

    Parameters
    ----------
    image_data_url : str
        Base64 data URL of the image.
    image_name : str
        Name of the image file.
    language : str, optional
        Language code for OCR (default is 'auto').

    Returns
    -------
    str
        Extracted text.
    """
    try:
        if "," in image_data_url:
            base64_data = image_data_url.split(",", 1)[1]
        else:
            base64_data = image_data_url

        image_bytes = base64.b64decode(base64_data)

        payload = {
            "isOverlayRequired": False,
            "apikey": settings.OCR_API_KEY,
            "language": language,
            "OCREngine": 2,
        }

        async with httpx.AsyncClient(timeout=30.0) as client:
            files = {"file": (image_name, BytesIO(image_bytes), "image/jpeg")}
            response = await client.post("https://api.ocr.space/parse/image", files=files, data=payload)

            if response.status_code == 200:
                result = response.json()

                if result.get("IsErroredOnProcessing", False):
                    error_msg = result.get("ErrorMessage", ["Unknown error"])[0]
                    ocr_logger.error(f"ocr.api_error image={image_name} error={error_msg}")
                    return f"[Ошибка OCR: {error_msg}]"

                parsed_results = result.get("ParsedResults", [])
                if parsed_results:
                    parsed_text = parsed_results[0].get("ParsedText", "")
                    ocr_logger.info(f"ocr.success image={image_name} text_len={len(parsed_text)}")
                    return parsed_text.strip()
                ocr_logger.warning(f"ocr.no_text image={image_name}")
                return "[Текст не обнаружен на изображении]"
            ocr_logger.error(f"ocr.http_error image={image_name} status={response.status_code}")
            return f"[Ошибка при обращении к OCR сервису: {response.status_code}]"

    except Exception as e:
        ocr_logger.error(f"ocr.process_error image={image_name} error={e}", exc_info=True)
        return "[Ошибка при обработке изображения]"


async def process_multiple_images_ocr(images: list[FileAttachment], language: str = "eng") -> list[str]:
    """Process multiple images in parallel for OCR.

    Parameters
    ----------
    images : list[FileAttachment]
        List of images to process.
    language : str, optional
        Language code for OCR (default is 'eng').

    Returns
    -------
    list[str]
        List of extracted texts.
    """
    if not images:
        return []

    ocr_logger.info(f"ocr.batch_start count={len(images)}")

    tasks = [process_image_ocr(img.dataUrl, img.name, language) for img in images]
    results = await asyncio.gather(*tasks, return_exceptions=True)

    ocr_texts = []
    error_count = 0
    for i, result in enumerate(results):
        if isinstance(result, Exception):
            ocr_logger.error(f"ocr.batch_error image={images[i].name} error={result}")
            ocr_texts.append(f"[Ошибка при распознавании изображения '{images[i].name}']")
            error_count += 1
        else:
            ocr_texts.append(result)

    ocr_logger.info(f"ocr.batch_complete count={len(images)} success={len(images) - error_count} errors={error_count}")
    return ocr_texts


# ── Helpers ────────────────────────────────────────────────────────────────────


def decode_base64_file(data_url: str) -> bytes:
    """Decode base64 data URL to bytes."""
    if "," in data_url:
        base64_data = data_url.split(",", 1)[1]
    else:
        base64_data = data_url
    return base64.b64decode(base64_data)


_AUDIO_EXTENSIONS = {".mp3", ".wav", ".m4a", ".ogg", ".flac", ".webm"}

_whisper_model = None
_whisper_lock = asyncio.Lock()


def _get_whisper_model() -> object:
    global _whisper_model  # type: ignore[misc]
    if _whisper_model is None:
        from faster_whisper import WhisperModel

        _whisper_model = WhisperModel("tiny", device="cpu", compute_type="int8", cpu_threads=1)
    return _whisper_model


def detect_attachment_type(filename: str, data_url: str) -> AttachmentType:
    """Detect attachment type from filename or MIME type in data URL."""
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        return AttachmentType.PDF
    if filename_lower.endswith(".docx"):
        return AttachmentType.DOCX
    if filename_lower.endswith(".doc"):
        return AttachmentType.DOCX  # Will show error later for old .doc format
    if any(filename_lower.endswith(ext) for ext in _AUDIO_EXTENSIONS):
        return AttachmentType.AUDIO
    if data_url.startswith("data:application/pdf"):
        return AttachmentType.PDF
    if "wordprocessingml" in data_url or "msword" in data_url:
        return AttachmentType.DOCX
    if data_url.startswith("data:audio/"):
        return AttachmentType.AUDIO
    return AttachmentType.IMAGE


# ── Document extraction ────────────────────────────────────────────────────────


async def extract_pdf_text(file_bytes: bytes, filename: str) -> tuple[str, list[bytes]]:
    """Extract text and images from PDF file bytes.

    Returns
    -------
    tuple[str, list[bytes]]
        (extracted_text, list_of_image_bytes)
    """
    try:

        def _extract() -> tuple[str, list]:
            try:
                reader = PdfReader(BytesIO(file_bytes))
            except PdfReadError as e:
                if "password" in str(e).lower() or "encrypt" in str(e).lower():
                    return "[Ошибка: PDF файл защищен паролем]", []
                raise

            if len(reader.pages) == 0:
                return "[PDF файл пуст]", []

            text_parts = []
            images = []

            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(page_text)

                if len(images) < MAX_IMAGES_PER_DOCUMENT:
                    try:
                        for image_obj in page.images:
                            if len(images) >= MAX_IMAGES_PER_DOCUMENT:
                                break
                            try:
                                img_data = image_obj.data
                                img = Image.open(BytesIO(img_data))
                                if (img.width >= MIN_IMAGE_DIMENSION and img.height >= MIN_IMAGE_DIMENSION):
                                    images.append((img_data, page_num + 1))
                            except Exception:  # noqa: S112
                                continue
                    except Exception:  # noqa: S110
                        pass

            return "\n\n".join(text_parts), images

        text, images = await asyncio.get_event_loop().run_in_executor(None, _extract)
        doc_logger.info(f"pdf.extract filename={filename} text_len={len(text)} images={len(images)}")
        return text.strip() if text else "", images

    except Exception as e:
        doc_logger.error(f"pdf.extract_error filename={filename} error={e}", exc_info=True)
        return f"[Ошибка при извлечении текста из PDF: {e!s}]", []


async def extract_docx_text(file_bytes: bytes, filename: str) -> tuple[str, list[bytes]]:
    """Extract text and images from DOCX file bytes.

    Returns
    -------
    tuple[str, list[bytes]]
        (extracted_text, list_of_image_bytes)
    """
    if filename.lower().endswith(".doc") and not filename.lower().endswith(".docx"):
        return "[Ошибка: формат .doc не поддерживается, используйте .docx]", []

    try:

        def _extract() -> tuple[str, list]:
            doc = Document(BytesIO(file_bytes))
            text_parts = []
            images = []

            for para in doc.paragraphs:
                para_text = para.text.strip()
                if para_text:
                    text_parts.append(para_text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            try:
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        if len(images) >= MAX_IMAGES_PER_DOCUMENT:
                            break
                        try:
                            img_data = rel.target_part.blob
                            img = Image.open(BytesIO(img_data))
                            if (img.width >= MIN_IMAGE_DIMENSION and img.height >= MIN_IMAGE_DIMENSION):
                                images.append((img_data, None))
                        except Exception:  # noqa: S112
                            continue
            except Exception:  # noqa: S110
                pass

            return "\n\n".join(text_parts), images

        text, images = await asyncio.get_event_loop().run_in_executor(None, _extract)
        doc_logger.info(f"docx.extract filename={filename} text_len={len(text)} images={len(images)}")
        return text.strip() if text else "", images

    except Exception as e:
        doc_logger.error(f"docx.extract_error filename={filename} error={e}", exc_info=True)
        return f"[Ошибка при извлечении текста из DOCX: {e!s}]", []


async def ocr_document_images(images: list[tuple[bytes, int | None]], filename: str) -> list[str]:
    """Run OCR on extracted document images.

    Parameters
    ----------
    images : list[tuple[bytes, int | None]]
        List of (image_bytes, page_number) tuples.
    filename : str
        Document filename for logging.

    Returns
    -------
    list[str]
        List of OCR texts.
    """
    if not images:
        return []

    doc_logger.info(f"ocr.images_start filename={filename} count={len(images)}")

    async def _ocr_single(i: int, img_data: bytes, page_num: int | None) -> str | None:
        try:
            img = Image.open(BytesIO(img_data))
            img_format = img.format or "PNG"

            buffer = BytesIO()
            img.save(buffer, format=img_format)
            img_base64 = base64.b64encode(buffer.getvalue()).decode("utf-8")
            data_url = f"data:image/{img_format.lower()};base64,{img_base64}"

            page_suffix = f"_page_{page_num}" if page_num else ""
            ext = img_format.lower()
            if ext == "jpeg":
                ext = "jpg"
            img_name = f"{filename}_img_{i+1}{page_suffix}.{ext}"
            ocr_text = await process_image_ocr(data_url, img_name, settings.OCR_LANGUAGE)

            if (ocr_text and not ocr_text.startswith("[Ошибка") and not ocr_text.startswith("[Текст не")):
                location = f"стр. {page_num}" if page_num else f"изображение {i+1}"
                return f"[Текст из изображения в документе ({location})]:\n{ocr_text}"
        except Exception as e:
            doc_logger.warning(f"ocr.image_error filename={filename} image={i+1} error={e}")
        return None

    tasks = [_ocr_single(i, img_data, page_num) for i, (img_data, page_num) in enumerate(images)]
    results = await asyncio.gather(*tasks)
    ocr_results = [r for r in results if r]

    doc_logger.info(f"ocr.images_complete filename={filename} success={len(ocr_results)}")
    return ocr_results


# ── Audio ──────────────────────────────────────────────────────────────────────


async def transcribe_audio(file_bytes: bytes, filename: str) -> tuple[str, bool]:
    """Transcribe audio file using Faster-Whisper (tiny model, int8, single thread).

    Returns
    -------
    tuple[str, bool]
        (transcription_text, is_large)
    """
    suffix = Path(filename).suffix or ".audio"
    try:
        async with _whisper_lock:
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            try:
                model = _get_whisper_model()

                def _transcribe() -> tuple[str, object]:
                    segs, inf = model.transcribe(tmp_path, beam_size=1)
                    return "".join(seg.text for seg in segs).strip(), inf

                text, info = await asyncio.wait_for(
                    asyncio.get_event_loop().run_in_executor(None, _transcribe),
                    timeout=settings.WHISPER_TIMEOUT,
                )
            finally:
                os.unlink(tmp_path)

        doc_logger.info(f"audio.transcribe filename={filename} text_len={len(text)} language={info.language}")
        is_large = len(text) > LARGE_DOCUMENT_THRESHOLD
        return text or "[Аудио не содержит распознанной речи]", is_large

    except TimeoutError:
        doc_logger.error(f"audio.transcribe_timeout filename={filename}")
        return "[Ошибка: превышено время обработки аудио]", False

    except Exception as e:
        doc_logger.error(f"audio.transcribe_error filename={filename} error={e}", exc_info=True)
        return f"[Ошибка при транскрибации аудио: {e!s}]", False


# ── High-level pipeline ────────────────────────────────────────────────────────


async def process_document_with_ocr(attachment: FileAttachment) -> tuple[str, bool]:
    """Process a document attachment, extracting text and running OCR on embedded images.

    Parameters
    ----------
    attachment : FileAttachment
        The document attachment.

    Returns
    -------
    tuple[str, bool]
        (combined_text, is_large_document)
    """
    file_bytes = decode_base64_file(attachment.dataUrl)
    filename = attachment.name

    att_type = attachment.type
    if att_type is None:
        att_type = detect_attachment_type(filename, attachment.dataUrl)

    if att_type == AttachmentType.PDF:
        text, images = await extract_pdf_text(file_bytes, filename)
    elif att_type == AttachmentType.DOCX:
        text, images = await extract_docx_text(file_bytes, filename)
    else:
        return None, False  # Not a document

    if text.startswith("[Ошибка"):
        return text, False

    ocr_texts = await ocr_document_images(images, filename)

    result_parts = []
    if text:
        result_parts.append(text)
    if ocr_texts:
        result_parts.extend(ocr_texts)

    if not result_parts:
        return "[Документ пуст или не содержит извлекаемого текста]", False

    combined_text = "\n\n".join(result_parts)
    is_large = len(combined_text) > LARGE_DOCUMENT_THRESHOLD

    return combined_text, is_large


async def process_multiple_documents(
    attachments: list[FileAttachment],
) -> tuple[list[tuple[FileAttachment, str | None]], bool]:
    """Process multiple attachments, extracting text from documents.

    Parameters
    ----------
    attachments : list[FileAttachment]
        List of attachments to process.

    Returns
    -------
    tuple[list[tuple[FileAttachment, str | None]], bool]
        (results_list, has_large_document) where results_list is a list of
        (attachment, extracted_text_or_None) tuples.
    """
    if not attachments:
        return [], False

    doc_logger.info(f"doc.batch_start count={len(attachments)}")

    results = []
    has_large_document = False

    for att in attachments:
        att_type = detect_attachment_type(att.name, att.dataUrl)

        if att_type in (AttachmentType.PDF, AttachmentType.DOCX):
            text, is_large = await process_document_with_ocr(att)
            results.append((att, text))
            if is_large:
                has_large_document = True
        elif att_type == AttachmentType.AUDIO:
            file_bytes = decode_base64_file(att.dataUrl)
            text, is_large = await transcribe_audio(file_bytes, att.name)
            results.append((att, text))
            if is_large:
                has_large_document = True
        else:
            results.append((att, None))  # Image — will be processed by OCR separately

    doc_count = sum(1 for _, text in results if text is not None)
    doc_logger.info(f"doc.batch_complete total={len(attachments)} documents={doc_count} large={has_large_document}")

    return results, has_large_document
